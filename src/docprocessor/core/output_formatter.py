"""Output formatting for synthesized books in various formats."""

import platform
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from config.settings import settings
from docprocessor.models import Chapter
from docprocessor.utils.language_detector import LanguageDetector
from docprocessor.utils.logger import get_logger

logger = get_logger(__name__)


class OutputFormatter:
    """Formats synthesized books for various output formats."""

    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize output formatter.

        Args:
            output_dir: Directory for output files (default: from settings)
        """
        self.output_dir = output_dir or settings.output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

        logger.info(f"OutputFormatter initialized with output_dir: {self.output_dir}")

    def export_markdown(
        self,
        chapters: List[Chapter],
        title: str = "Synthesized Document",
        author: Optional[str] = None,
        output_filename: Optional[str] = None,
    ) -> Path:
        """
        Export chapters to Markdown format.

        Args:
            chapters: List of Chapter objects
            title: Book title
            author: Optional author name
            output_filename: Optional filename (auto-generated if not provided)

        Returns:
            Path to output file
        """
        logger.info(f"Exporting {len(chapters)} chapters to Markdown")

        # Generate filename
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{title.replace(' ', '_')}_{timestamp}.md"

        output_path = self.output_dir / output_filename

        # Detect language from chapter content
        sample_text = "\n".join(ch.content[:500] for ch in chapters[:3])
        language = LanguageDetector.detect_language(sample_text)
        labels = LanguageDetector.get_labels(language)

        logger.debug(f"Detected language: {language}")

        # Build markdown content
        md_lines = []

        # Title page
        md_lines.append(f"# {title}\n")
        if author:
            md_lines.append(f"**{labels['by']}**: {author}\n")
        md_lines.append(f"**{labels['generated']}**: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
        md_lines.append("---\n")

        # Table of contents
        md_lines.append(f"## {labels['table_of_contents']}\n")
        for chapter in chapters:
            md_lines.append(
                f"{chapter.chapter_number}. [{chapter.title}](#chapter-{chapter.chapter_number})\n"
            )
        md_lines.append("\n---\n")

        # Chapters
        for chapter in chapters:
            md_lines.append(
                f"\n## {labels['chapter']} {chapter.chapter_number}: {chapter.title} {{#chapter-{chapter.chapter_number}}}\n"
            )
            md_lines.append(chapter.content)
            md_lines.append("\n")

            # Add citations if present
            if chapter.citations:
                md_lines.append("\n### References\n")
                for i, citation in enumerate(chapter.citations, 1):
                    doc_id = citation.get("document_id", "unknown")
                    page = citation.get("page", "unknown")
                    md_lines.append(f"- [{i}] Document {doc_id}, Page {page}\n")

            md_lines.append("\n---\n")

        # Write to file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))

        logger.info(f"Markdown exported to: {output_path}")
        return output_path

    def export_docx(
        self,
        chapters: List[Chapter],
        title: str = "Synthesized Document",
        author: Optional[str] = None,
        output_filename: Optional[str] = None,
    ) -> Path:
        """
        Export chapters to DOCX format.

        Args:
            chapters: List of Chapter objects
            title: Book title
            author: Optional author name
            output_filename: Optional filename (auto-generated if not provided)

        Returns:
            Path to output file
        """
        logger.info(f"Exporting {len(chapters)} chapters to DOCX")

        # Generate filename
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{title.replace(' ', '_')}_{timestamp}.docx"

        output_path = self.output_dir / output_filename

        # Create document
        doc = DocxDocument()

        # Title page
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if author:
            author_para = doc.add_paragraph()
            author_run = author_para.add_run(f"By {author}")
            author_run.font.size = Pt(14)
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        date_run.font.size = Pt(12)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Table of contents
        toc_heading = doc.add_heading("Table of Contents", level=1)
        for chapter in chapters:
            toc_para = doc.add_paragraph(f"Chapter {chapter.chapter_number}: {chapter.title}")
            toc_para.style = "List Number"

        doc.add_page_break()

        # Chapters
        for chapter in chapters:
            # Chapter heading
            doc.add_heading(f"Chapter {chapter.chapter_number}: {chapter.title}", level=1)

            # Chapter content (parse and add paragraphs)
            self._add_formatted_content(doc, chapter.content)

            # Citations
            if chapter.citations:
                doc.add_heading("References", level=2)
                for i, citation in enumerate(chapter.citations, 1):
                    doc_id = citation.get("document_id", "unknown")
                    page = citation.get("page", "unknown")
                    cite_para = doc.add_paragraph(f"[{i}] Document {doc_id}, Page {page}")
                    cite_para.style = "List Bullet"

            # Page break after each chapter
            doc.add_page_break()

        # Save document
        doc.save(str(output_path))

        logger.info(f"DOCX exported to: {output_path}")
        return output_path

    def _add_formatted_content(self, doc: DocxDocument, content: str) -> None:
        """
        Add formatted content to DOCX document.

        Handles headings, paragraphs, and basic formatting.

        Args:
            doc: DOCX document object
            content: Content text to add
        """
        lines = content.split("\n")
        current_paragraph = []

        for line in lines:
            line = line.strip()

            # Skip empty lines between paragraphs
            if not line:
                if current_paragraph:
                    doc.add_paragraph(" ".join(current_paragraph))
                    current_paragraph = []
                continue

            # Check for headings (markdown style)
            if line.startswith("###"):
                if current_paragraph:
                    doc.add_paragraph(" ".join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line.replace("###", "").strip(), level=3)
            elif line.startswith("##"):
                if current_paragraph:
                    doc.add_paragraph(" ".join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line.replace("##", "").strip(), level=2)
            elif line.startswith("#"):
                if current_paragraph:
                    doc.add_paragraph(" ".join(current_paragraph))
                    current_paragraph = []
                doc.add_heading(line.replace("#", "").strip(), level=1)
            else:
                # Regular text - accumulate into paragraph
                current_paragraph.append(line)

        # Add final paragraph if any
        if current_paragraph:
            doc.add_paragraph(" ".join(current_paragraph))

    def export_plain_text(
        self,
        chapters: List[Chapter],
        title: str = "Synthesized Document",
        author: Optional[str] = None,
        output_filename: Optional[str] = None,
    ) -> Path:
        """
        Export chapters to plain text format.

        Args:
            chapters: List of Chapter objects
            title: Book title
            author: Optional author name
            output_filename: Optional filename (auto-generated if not provided)

        Returns:
            Path to output file
        """
        logger.info(f"Exporting {len(chapters)} chapters to plain text")

        # Generate filename
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{title.replace(' ', '_')}_{timestamp}.txt"

        output_path = self.output_dir / output_filename

        # Build text content
        text_lines = []

        # Title
        text_lines.append("=" * 80)
        text_lines.append(title.center(80))
        text_lines.append("=" * 80)
        text_lines.append("")

        if author:
            text_lines.append(f"Author: {author}")
            text_lines.append("")

        text_lines.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        text_lines.append("")
        text_lines.append("=" * 80)
        text_lines.append("")

        # Table of contents
        text_lines.append("TABLE OF CONTENTS")
        text_lines.append("-" * 80)
        for chapter in chapters:
            text_lines.append(f"  {chapter.chapter_number}. {chapter.title}")
        text_lines.append("")
        text_lines.append("=" * 80)
        text_lines.append("")

        # Chapters
        for chapter in chapters:
            text_lines.append("")
            text_lines.append("=" * 80)
            text_lines.append(f"CHAPTER {chapter.chapter_number}: {chapter.title}".center(80))
            text_lines.append("=" * 80)
            text_lines.append("")
            text_lines.append(chapter.content)
            text_lines.append("")

            # Citations
            if chapter.citations:
                text_lines.append("-" * 80)
                text_lines.append("REFERENCES:")
                for i, citation in enumerate(chapter.citations, 1):
                    doc_id = citation.get("document_id", "unknown")
                    page = citation.get("page", "unknown")
                    text_lines.append(f"  [{i}] Document {doc_id}, Page {page}")
                text_lines.append("")

        # Write to file
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(text_lines))

        logger.info(f"Plain text exported to: {output_path}")
        return output_path

    def _find_pandoc(self) -> Optional[str]:
        """Find pandoc executable in common locations.

        Returns:
            Path to pandoc executable or None if not found
        """
        # Try shutil.which first (checks PATH)
        pandoc_path = shutil.which("pandoc")
        if pandoc_path:
            return pandoc_path

        # Check common installation locations
        common_paths = [
            "/usr/local/bin/pandoc",
            "/opt/homebrew/bin/pandoc",
            "/usr/bin/pandoc",
            "C:\\Program Files\\Pandoc\\pandoc.exe",
            "C:\\Program Files (x86)\\Pandoc\\pandoc.exe",
        ]

        for path in common_paths:
            if Path(path).exists():
                return path

        return None

    def _find_xelatex(self) -> Optional[str]:
        """Find xelatex executable in common locations.

        Returns:
            Path to xelatex executable or None if not found
        """
        # Try shutil.which first (checks PATH)
        xelatex_path = shutil.which("xelatex")
        if xelatex_path:
            return xelatex_path

        # Check common TeX Live installation locations
        common_paths = [
            "/usr/local/texlive/2025/bin/universal-darwin/xelatex",
            "/usr/local/texlive/2024/bin/universal-darwin/xelatex",
            "/usr/local/texlive/2023/bin/universal-darwin/xelatex",
            "/Library/TeX/texbin/xelatex",
            "/usr/texbin/xelatex",
        ]

        for path in common_paths:
            if Path(path).exists():
                return path

        return None

    def _check_pdf_requirements(self) -> tuple[bool, str]:
        """Check if PDF generation requirements are met.

        Returns:
            Tuple of (requirements_met, error_message)
        """
        # Check pandoc
        pandoc_path = self._find_pandoc()
        if not pandoc_path:
            # Platform-specific error messages
            system = platform.system()
            if system == "Darwin":  # macOS
                msg = "Pandoc not installed. Install with: brew install pandoc or visit https://pandoc.org/installing.html"
            elif system == "Windows":
                msg = "Pandoc not installed. Install from: https://pandoc.org/installing.html or run dependency checker"
            else:  # Linux
                msg = "Pandoc not installed. Install with: sudo apt install pandoc (Ubuntu) or visit https://pandoc.org/installing.html"
            return (False, msg)

        # Check xelatex (PDF engine)
        xelatex_path = self._find_xelatex()
        if not xelatex_path:
            # Platform-specific error messages
            system = platform.system()
            if system == "Darwin":  # macOS
                msg = "XeLaTeX not installed. Install MacTeX: brew install --cask mactex or visit https://www.tug.org/mactex/"
            elif system == "Windows":
                msg = "XeLaTeX not installed. Install MiKTeX from: https://miktex.org/download or TeX Live from: https://www.tug.org/texlive/"
            else:  # Linux
                msg = "XeLaTeX not installed. Install with: sudo apt install texlive-xetex (Ubuntu) or visit https://www.tug.org/texlive/"
            return (False, msg)

        return True, ""

    def export_pdf(
        self,
        chapters: List[Chapter],
        title: str = "Synthesized Document",
        author: Optional[str] = None,
        output_filename: Optional[str] = None,
        template: str = "academic",
    ) -> Path:
        """
        Export chapters to PDF format using pandoc.

        Args:
            chapters: List of Chapter objects
            title: Book title
            author: Optional author name
            output_filename: Optional filename (auto-generated if not provided)
            template: PDF template style (academic, professional, simple)

        Returns:
            Path to output file

        Raises:
            Exception: If pandoc is not installed or PDF generation fails
        """
        logger.info(f"Exporting {len(chapters)} chapters to PDF")

        # Check PDF generation requirements
        requirements_met, error_msg = self._check_pdf_requirements()
        if not requirements_met:
            raise Exception(error_msg)

        # Find pandoc and xelatex (we know they exist from the check above)
        pandoc_path = self._find_pandoc()
        xelatex_path = self._find_xelatex()

        # Generate filename
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{title.replace(' ', '_')}_{timestamp}.pdf"

        output_path = self.output_dir / output_filename

        # First generate markdown
        md_filename = output_filename.replace(".pdf", "_temp.md")
        md_path = self.export_markdown(chapters, title, author, md_filename)

        # Try to convert with pandoc
        try:
            logger.info(f"Using pandoc at: {pandoc_path}")
            logger.info(f"Using xelatex at: {xelatex_path}")

            if True:  # Always try if we found pandoc
                logger.info("Using pandoc for PDF generation")

                # Pandoc command with template options
                pandoc_args = [
                    pandoc_path,  # Use full path to pandoc
                    str(md_path),
                    "-o",
                    str(output_path),
                    f"--pdf-engine={xelatex_path}",  # Use full path to xelatex
                    "-V",
                    "geometry:margin=1in",
                    "-V",
                    f"title={title}",
                ]

                if author:
                    pandoc_args.extend(["-V", f"author={author}"])

                # Template-specific options
                if template == "academic":
                    pandoc_args.extend(
                        [
                            "-V",
                            "documentclass=report",
                            "-V",
                            "fontsize=12pt",
                            "--toc",
                            "--toc-depth=2",
                            "-V",
                            "colorlinks=true",
                        ]
                    )
                elif template == "professional":
                    pandoc_args.extend(
                        [
                            "-V",
                            "documentclass=article",
                            "-V",
                            "fontsize=11pt",
                            "--toc",
                        ]
                    )
                else:  # simple
                    pandoc_args.extend(
                        [
                            "-V",
                            "documentclass=article",
                            "-V",
                            "fontsize=12pt",
                        ]
                    )

                # Run pandoc (hide console window on Windows)
                startupinfo = None
                if platform.system() == "Windows":
                    startupinfo = subprocess.STARTUPINFO()
                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                    startupinfo.wShowWindow = subprocess.SW_HIDE

                result = subprocess.run(
                    pandoc_args,
                    capture_output=True,
                    text=True,
                    timeout=120,
                    startupinfo=startupinfo,
                )

                if result.returncode == 0:
                    # Clean up temp markdown
                    try:
                        md_path.unlink()
                    except Exception:
                        pass  # Ignore cleanup errors
                    logger.info(f"PDF exported to: {output_path}")
                    return output_path
                else:
                    # Clean up temp markdown on failure
                    try:
                        md_path.unlink()
                    except Exception:
                        pass
                    logger.error(f"Pandoc error: {result.stderr}")
                    raise Exception(f"Pandoc failed: {result.stderr}")

        except subprocess.TimeoutExpired:
            # Clean up temp markdown on timeout
            try:
                md_path.unlink()
            except Exception:
                pass
            logger.error("Pandoc timeout")
            raise Exception("PDF generation timed out")
        except Exception as e:
            # Clean up temp markdown on any error
            try:
                md_path.unlink()
            except Exception:
                pass
            logger.error(f"Error generating PDF: {e}")
            raise

    def export_all_formats(
        self,
        chapters: List[Chapter],
        title: str = "Synthesized Document",
        author: Optional[str] = None,
        include_pdf: bool = False,
        pdf_template: str = "academic",
    ) -> Dict[str, Path]:
        """
        Export to all supported formats.

        Args:
            chapters: List of Chapter objects
            title: Book title
            author: Optional author name
            include_pdf: Whether to include PDF export
            pdf_template: PDF template style

        Returns:
            Dictionary mapping format names to output paths
        """
        logger.info("Exporting to all formats")

        outputs = {}

        try:
            outputs["markdown"] = self.export_markdown(chapters, title, author)
        except Exception as e:
            logger.error(f"Error exporting to Markdown: {e}")

        try:
            outputs["docx"] = self.export_docx(chapters, title, author)
        except Exception as e:
            logger.error(f"Error exporting to DOCX: {e}")

        try:
            outputs["text"] = self.export_plain_text(chapters, title, author)
        except Exception as e:
            logger.error(f"Error exporting to plain text: {e}")

        if include_pdf:
            try:
                outputs["pdf"] = self.export_pdf(chapters, title, author, template=pdf_template)
            except Exception as e:
                logger.error(f"Error exporting to PDF: {e}")
                # Platform-specific warning message
                system = platform.system()
                if system == "Darwin":
                    logger.warning("PDF export failed. Install pandoc: brew install pandoc")
                elif system == "Windows":
                    logger.warning("PDF export failed. Install pandoc from: https://pandoc.org/installing.html")
                else:
                    logger.warning("PDF export failed. Install pandoc: sudo apt install pandoc")

        logger.info(f"Exported to {len(outputs)} formats")
        return outputs
