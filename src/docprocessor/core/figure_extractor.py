"""Figure and statistic extraction from documents."""

import logging
import re
from pathlib import Path
from typing import List, Optional, Tuple

from docprocessor.models.extracted_figure import ExtractedFigure, FigureExtractionResult, FigureType

logger = logging.getLogger(__name__)


class FigureExtractor:
    """Extract figures, statistics, dates, and currency from documents."""

    def __init__(self):
        """Initialize the figure extractor."""
        # Regex patterns for different figure types
        self.patterns = {
            # Currency patterns (€, $, TND, EUR, USD, etc.)
            "currency": [
                r"(€|EUR)\s*(\d+(?:[,.\s]\d+)*(?:[,.]?\d+)?)\s*(milliards?|millions?|milliers?|mille|k|M|B)?",
                r"(\d+(?:[,.\s]\d+)*(?:[,.]?\d+)?)\s*(€|EUR|euros?)",
                r"\$\s*(\d+(?:[,.\s]\d+)*(?:[,.]?\d+)?)\s*(billions?|millions?|thousands?|k|M|B)?",
                r"(\d+(?:[,.\s]\d+)*(?:[,.]?\d+)?)\s*TND",
                r"(\d+(?:[,.\s]\d+)*(?:[,.]?\d+)?)\s*dinars?",
            ],
            # Percentage patterns
            "percentage": [
                r"(\d+(?:[,.]\d+)?)\s*%",
                r"(\d+(?:[,.]\d+)?)\s*pour\s*cent",
                r"(\d+(?:[,.]\d+)?)\s*percent",
            ],
            # Date patterns (years, full dates)
            "date": [
                r"\b(20\d{2})\b",  # Years 2000-2099
                r"\b(19\d{2})\b",  # Years 1900-1999
                r"\b(\d{1,2})[/-](\d{1,2})[/-](20\d{2})\b",  # DD/MM/YYYY or MM/DD/YYYY
                r"\b(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre)\s+(20\d{2})\b",
                r"\b(january|february|march|april|may|june|july|august|september|october|november|december)\s+(20\d{2})\b",
            ],
            # Range patterns (2020-2025, 2020–2025)
            "range": [
                r"(20\d{2})\s*[-–—]\s*(20\d{2})",  # Year ranges
                r"(\d+(?:[,.]\d+)?)\s*[-–—à]\s*(\d+(?:[,.]\d+)?)\s*%",  # Percentage ranges
            ],
            # Quantity with units
            "quantity": [
                r"(\d+(?:[,.\s]\d+)*)\s+(fonctionnaires?|employés?|personnes?|habitants?|citoyens?)",
                r"(\d+(?:[,.\s]\d+)*)\s+(tonnes?|kg|grammes?|litres?|m[²³]?|km)",
                r"(\d+(?:[,.\s]\d+)*)\s+(jours?|mois|années?|heures?)",
            ],
            # Generic numbers (last resort, captures standalone numbers)
            "number": [
                r"\b(\d+(?:[,.\s]\d+)*(?:[,.]\d+)?)\b",
            ],
        }

    def extract_from_text(
        self, text: str, page_number: Optional[int] = None, paragraph_number: Optional[int] = None
    ) -> List[ExtractedFigure]:
        """Extract figures from plain text.

        Args:
            text: Text to extract from
            page_number: Optional page number
            paragraph_number: Optional paragraph number

        Returns:
            List of extracted figures
        """
        figures = []

        # Split into sentences for context
        sentences = self._split_sentences(text)

        for sent_idx, sentence in enumerate(sentences):
            # Try each pattern type in order (most specific first)
            for pattern_type, patterns in self.patterns.items():
                for pattern in patterns:
                    matches = re.finditer(pattern, sentence, re.IGNORECASE)

                    for match in matches:
                        # Extract the full matched value
                        value = match.group(0)

                        # Determine figure type
                        if pattern_type == "currency":
                            figure_type = FigureType.CURRENCY
                        elif pattern_type == "percentage":
                            figure_type = FigureType.PERCENTAGE
                        elif pattern_type == "date":
                            figure_type = FigureType.DATE
                        elif pattern_type == "range":
                            figure_type = FigureType.RANGE
                        elif pattern_type == "quantity":
                            figure_type = FigureType.QUANTITY
                        else:
                            figure_type = FigureType.NUMBER

                        # Parse numeric value and metadata
                        numeric_value, unit, currency_code, year = self._parse_figure(
                            value, figure_type
                        )

                        # If year not in value, try to extract from sentence context
                        if not year:
                            year_in_context = re.search(r"\b(19|20)\d{2}\b", sentence)
                            if year_in_context:
                                year = int(year_in_context.group(0))

                        # Get surrounding context
                        start_pos = match.start()
                        end_pos = match.end()
                        surrounding_text = self._get_surrounding_text(sentence, start_pos, end_pos)

                        # Create ExtractedFigure
                        figure = ExtractedFigure(
                            value=value.strip(),
                            figure_type=figure_type,
                            page_number=page_number,
                            paragraph_number=paragraph_number,
                            sentence_index=sent_idx,
                            context_sentence=sentence.strip(),
                            context_paragraph=text[:500],  # First 500 chars as paragraph context
                            surrounding_text=surrounding_text,
                            numeric_value=numeric_value,
                            unit=unit,
                            currency_code=currency_code,
                            year=year,
                        )

                        figures.append(figure)

        # Remove duplicates (same value in same sentence)
        figures = self._remove_duplicates(figures)

        return figures

    def extract_from_document(self, document_path: Path) -> FigureExtractionResult:
        """Extract figures from a document file (DOCX, PDF, TXT).

        Args:
            document_path: Path to document

        Returns:
            FigureExtractionResult with all extracted figures
        """
        import time

        start_time = time.time()

        result = FigureExtractionResult(document_path=str(document_path))

        try:
            # Determine file type
            suffix = document_path.suffix.lower()

            if suffix == ".txt":
                figures = self._extract_from_txt(document_path)
                result.figures = figures

            elif suffix == ".docx":
                # Extract from DOCX (text + tables)
                text_figures = self._extract_from_docx_text(document_path)
                table_figures = self._extract_from_docx_tables(document_path)
                result.figures = text_figures + table_figures
                result.tables_parsed = len(self._get_docx_tables(document_path))

            elif suffix == ".pdf":
                # Extract from PDF (text + tables)
                text_figures = self._extract_from_pdf_text(document_path)
                # table_figures = self._extract_from_pdf_tables(document_path)  # TODO: Implement
                result.figures = text_figures
                result.tables_parsed = 0  # TODO: Count PDF tables

            else:
                result.errors.append(f"Unsupported file type: {suffix}")

        except Exception as e:
            logger.error(f"Error extracting figures from {document_path}: {e}")
            result.errors.append(str(e))

        # Calculate statistics
        result.total_figures = len(result.figures)
        result.extraction_time_seconds = time.time() - start_time

        # Update figures_by_type
        result.figures_by_type = {}
        for figure in result.figures:
            fig_type = figure.figure_type.value
            result.figures_by_type[fig_type] = result.figures_by_type.get(fig_type, 0) + 1

        return result

    def _extract_from_txt(self, file_path: Path) -> List[ExtractedFigure]:
        """Extract from plain text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return self.extract_from_text(text)

    def _extract_from_docx_text(self, file_path: Path) -> List[ExtractedFigure]:
        """Extract figures from DOCX text (excluding tables)."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed, cannot extract from DOCX")
            return []

        doc = Document(file_path)
        figures = []

        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                para_figures = self.extract_from_text(
                    paragraph.text,
                    page_number=None,  # DOCX doesn't have page numbers
                    paragraph_number=para_idx,
                )
                figures.extend(para_figures)

        return figures

    def _extract_from_docx_tables(self, file_path: Path) -> List[ExtractedFigure]:
        """Extract figures from DOCX tables."""
        try:
            from docx import Document
        except ImportError:
            return []

        doc = Document(file_path)
        figures = []

        for table_idx, table in enumerate(doc.tables):
            # Get headers (first row)
            headers = [cell.text.strip() for cell in table.rows[0].cells]

            # Process each row (skip header)
            for row_idx, row in enumerate(table.rows[1:], start=1):
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Extract figures from cell
                        cell_figures = self.extract_from_text(cell_text)

                        # Add table context
                        for figure in cell_figures:
                            figure.is_from_table = True
                            figure.table_index = table_idx
                            figure.table_row = row_idx
                            figure.table_column = col_idx
                            figure.table_column_header = (
                                headers[col_idx] if col_idx < len(headers) else None
                            )
                            # Row header is first cell in row
                            figure.table_row_header = (
                                row.cells[0].text.strip() if col_idx > 0 else None
                            )

                        figures.extend(cell_figures)

        return figures

    def _extract_from_pdf_text(self, file_path: Path) -> List[ExtractedFigure]:
        """Extract figures from PDF text."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed, cannot extract from PDF")
            return []

        doc = fitz.open(file_path)
        figures = []

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            if text.strip():
                page_figures = self.extract_from_text(text, page_number=page_num)
                figures.extend(page_figures)

        doc.close()
        return figures

    def _get_docx_tables(self, file_path: Path) -> list:
        """Get list of tables from DOCX."""
        try:
            from docx import Document

            doc = Document(file_path)
            return doc.tables
        except ImportError:
            return []

    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences (simple implementation)."""
        # Split on sentence boundaries, but not on decimal points
        # Use lookahead/lookbehind to avoid splitting on numbers
        sentences = re.split(r"(?<![0-9])[.!?]+\s+(?=[A-ZÀ-Ü])", text)
        # Also handle end of text
        if sentences:
            # If no splits occurred, return whole text as one sentence
            if len(sentences) == 1:
                return [text.strip()] if text.strip() else []
        return [s.strip() for s in sentences if s.strip()]

    def _get_surrounding_text(
        self, sentence: str, start_pos: int, end_pos: int, window: int = 50
    ) -> str:
        """Get text surrounding the match."""
        start = max(0, start_pos - window)
        end = min(len(sentence), end_pos + window)
        return sentence[start:end]

    def _parse_figure(
        self, value: str, figure_type: FigureType
    ) -> Tuple[Optional[float], Optional[str], Optional[str], Optional[int]]:
        """Parse figure to extract numeric value, unit, currency code, and year.

        Returns:
            (numeric_value, unit, currency_code, year)
        """
        numeric_value = None
        unit = None
        currency_code = None
        year = None

        # Extract year if present
        year_match = re.search(r"\b(19|20)\d{2}\b", value)
        if year_match:
            year = int(year_match.group(0))

        # Extract numeric value
        number_match = re.search(r"(\d+(?:[,.\s]\d+)*(?:[,.]?\d+)?)", value)
        if number_match:
            num_str = number_match.group(0)
            # Remove spaces
            num_str = num_str.replace(" ", "")

            # Detect format: European (1.234,56) vs US/UK (1,234.56)
            # If both . and , present, determine which is decimal separator
            if "," in num_str and "." in num_str:
                # Find last occurrence of each
                last_comma = num_str.rfind(",")
                last_dot = num_str.rfind(".")
                if last_comma > last_dot:
                    # European format: 1.234,56 → 1234.56
                    num_str = num_str.replace(".", "").replace(",", ".")
                else:
                    # US format: 1,234.56 → 1234.56
                    num_str = num_str.replace(",", "")
            elif "," in num_str:
                # Only comma: could be decimal (45,3) or thousands (1,234)
                # If there's only one comma and 1-2 digits after, it's decimal
                comma_parts = num_str.split(",")
                if len(comma_parts) == 2 and len(comma_parts[1]) <= 2:
                    # Decimal: 45,3 → 45.3
                    num_str = num_str.replace(",", ".")
                else:
                    # Thousands: 1,234 → 1234
                    num_str = num_str.replace(",", "")
            elif "." in num_str:
                # Only dot: could be decimal (45.3) or thousands (1.234)
                # If there's only one dot and 1-2 digits after, it's decimal
                dot_parts = num_str.split(".")
                if len(dot_parts) == 2 and len(dot_parts[1]) <= 2:
                    # Already decimal format
                    pass
                else:
                    # Thousands: 1.234.567 → 1234567
                    num_str = num_str.replace(".", "")

            try:
                numeric_value = float(num_str)
            except ValueError:
                pass

        # Detect currency
        if "€" in value or "EUR" in value or "euro" in value.lower():
            currency_code = "EUR"
            unit = "€"
        elif "$" in value or "USD" in value:
            currency_code = "USD"
            unit = "$"
        elif "TND" in value or "dinar" in value.lower():
            currency_code = "TND"
            unit = "TND"

        # Detect percentage
        if "%" in value or "pour cent" in value.lower() or "percent" in value.lower():
            unit = "%"

        # Detect multipliers (milliards, millions, etc.)
        if "milliard" in value.lower() or re.search(r"\bB\b", value):
            if numeric_value:
                numeric_value *= 1_000_000_000
            unit = (unit or "") + " milliards"
        elif "million" in value.lower() or re.search(r"\bM\b", value):
            if numeric_value:
                numeric_value *= 1_000_000
            unit = (unit or "") + " millions"
        elif (
            "millier" in value.lower()
            or "mille" in value.lower()
            or re.search(r"\bk\b", value, re.IGNORECASE)
        ):
            if numeric_value:
                numeric_value *= 1_000
            unit = (unit or "") + " milliers"

        return numeric_value, unit, currency_code, year

    def _remove_duplicates(self, figures: List[ExtractedFigure]) -> List[ExtractedFigure]:
        """Remove duplicate figures (same value in same sentence)."""
        seen = set()
        unique_figures = []

        for figure in figures:
            # Create unique key: (value, sentence_index)
            key = (figure.value, figure.sentence_index)
            if key not in seen:
                seen.add(key)
                unique_figures.append(figure)

        return unique_figures
