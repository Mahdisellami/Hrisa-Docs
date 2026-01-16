"""Document processing module for PDF text extraction and chunking."""

import re
from pathlib import Path

import pymupdf as fitz

from config.settings import settings
from docprocessor.models import Chunk, Document
from docprocessor.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentProcessor:
    """Handles PDF text extraction and document chunking."""

    def __init__(
        self,
        chunk_size: int = settings.chunk_size,
        chunk_overlap: int = settings.chunk_overlap,
    ):
        """
        Initialize document processor.

        Args:
            chunk_size: Maximum chunk size in tokens (approximate using words)
            chunk_overlap: Overlap between chunks in tokens
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(
            f"DocumentProcessor initialized with chunk_size={chunk_size}, "
            f"chunk_overlap={chunk_overlap}"
        )

    def extract_text_from_pdf(self, pdf_path: Path) -> Document:
        """
        Extract text and metadata from a PDF file.

        Args:
            pdf_path: Path to the PDF file

        Returns:
            Document object with extracted content and metadata

        Raises:
            ValueError: If file is not a valid PDF
            FileNotFoundError: If file doesn't exist
        """
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")

        if pdf_path.suffix.lower() != ".pdf":
            raise ValueError(f"File must be a PDF: {pdf_path}")

        logger.info(f"Extracting text from PDF: {pdf_path}")

        try:
            doc = fitz.open(pdf_path)

            title = doc.metadata.get("title", pdf_path.stem)
            author = doc.metadata.get("author", None)
            page_count = len(doc)
            file_size = pdf_path.stat().st_size

            full_text = ""
            for page_num, page in enumerate(doc, start=1):
                text = page.get_text()
                full_text += f"\n--- Page {page_num} ---\n{text}"

            doc.close()

            document = Document(
                file_path=pdf_path,
                title=title,
                author=author,
                page_count=page_count,
                file_size=file_size,
                text_content=full_text,
                processed=False,
            )

            logger.info(
                f"Successfully extracted text from {pdf_path.name}: "
                f"{page_count} pages, {len(full_text)} characters"
            )
            return document

        except fitz.FileDataError as e:
            logger.error(f"Invalid PDF file {pdf_path}: {e}")
            raise ValueError(f"Invalid PDF file: {e}")
        except Exception as e:
            logger.error(f"Error extracting text from {pdf_path}: {e}")
            raise

    def extract_text_from_txt(self, txt_path: Path) -> Document:
        """
        Extract text from a plain text file.

        Args:
            txt_path: Path to the text file

        Returns:
            Document object with extracted content

        Raises:
            FileNotFoundError: If file doesn't exist
        """
        if not txt_path.exists():
            raise FileNotFoundError(f"Text file not found: {txt_path}")

        logger.info(f"Extracting text from TXT: {txt_path}")

        try:
            with open(txt_path, "r", encoding="utf-8") as f:
                full_text = f.read()

            title = txt_path.stem
            file_size = txt_path.stat().st_size

            document = Document(
                file_path=txt_path,
                title=title,
                author=None,
                page_count=1,  # Text files don't have pages
                file_size=file_size,
                text_content=full_text,
                processed=False,
            )

            logger.info(
                f"Successfully extracted text from {txt_path.name}: " f"{len(full_text)} characters"
            )
            return document

        except Exception as e:
            logger.error(f"Error extracting text from {txt_path}: {e}")
            raise

    def extract_text_from_docx(self, docx_path: Path) -> Document:
        """
        Extract text from a DOCX file.

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Document object with extracted content

        Raises:
            FileNotFoundError: If file doesn't exist
        """
        if not docx_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")

        logger.info(f"Extracting text from DOCX: {docx_path}")

        try:
            # Try to import python-docx
            try:
                from docx import Document as DocxDocument
            except ImportError:
                logger.warning("python-docx not installed, reading DOCX as plain text")
                # Fallback: try to read as text (won't work well but better than nothing)
                with open(docx_path, "rb") as f:
                    content = f.read()
                    # Try to extract some text (very basic)
                    full_text = str(content)
                    title = docx_path.stem
                    file_size = docx_path.stat().st_size

                    document = Document(
                        file_path=docx_path,
                        title=title,
                        author=None,
                        page_count=1,
                        file_size=file_size,
                        text_content=full_text[:10000],  # Limit to reasonable size
                        processed=False,
                    )
                    return document

            # Use python-docx to extract text
            docx_doc = DocxDocument(docx_path)

            # Extract all paragraphs
            paragraphs = []
            for para in docx_doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            full_text = "\n\n".join(paragraphs)

            # Extract metadata
            core_props = docx_doc.core_properties
            title = core_props.title or docx_path.stem
            author = core_props.author
            file_size = docx_path.stat().st_size

            document = Document(
                file_path=docx_path,
                title=title,
                author=author,
                page_count=1,  # DOCX doesn't have fixed pages
                file_size=file_size,
                text_content=full_text,
                processed=False,
            )

            logger.info(
                f"Successfully extracted text from {docx_path.name}: "
                f"{len(paragraphs)} paragraphs, {len(full_text)} characters"
            )
            return document

        except Exception as e:
            logger.error(f"Error extracting text from {docx_path}: {e}")
            raise

    def chunk_document(self, document: Document) -> list[Chunk]:
        """
        Chunk document text into smaller semantic units.

        Uses a paragraph-based approach with size limits.

        Args:
            document: Document object with text content

        Returns:
            List of Chunk objects

        Raises:
            ValueError: If document has no text content
        """
        if not document.text_content:
            raise ValueError("Document has no text content to chunk")

        logger.info(f"Chunking document: {document.title}")

        paragraphs = self._split_into_paragraphs(document.text_content)
        chunks = []
        current_chunk_text = ""
        current_page = 1
        chunk_index = 0
        start_char = 0

        for paragraph in paragraphs:
            # Extract page number if present, but keep the text
            page_match = re.search(r"--- Page (\d+) ---", paragraph)
            if page_match:
                current_page = int(page_match.group(1))
                # Remove the page marker from the paragraph text
                paragraph = re.sub(r"--- Page \d+ ---\s*", "", paragraph)

            paragraph = paragraph.strip()
            if not paragraph:
                continue

            estimated_tokens = self._estimate_tokens(paragraph)

            if (
                self._estimate_tokens(current_chunk_text) + estimated_tokens > self.chunk_size
                and current_chunk_text
            ):
                chunk = self._create_chunk(
                    document_id=document.id,
                    text=current_chunk_text,
                    page_number=current_page,
                    chunk_index=chunk_index,
                    start_char=start_char,
                )
                chunks.append(chunk)
                chunk_index += 1

                overlap_text = self._get_overlap_text(current_chunk_text)
                start_char += len(current_chunk_text) - len(overlap_text)
                current_chunk_text = overlap_text + " " + paragraph
            else:
                if current_chunk_text:
                    current_chunk_text += " " + paragraph
                else:
                    current_chunk_text = paragraph

        if current_chunk_text.strip():
            chunk = self._create_chunk(
                document_id=document.id,
                text=current_chunk_text,
                page_number=current_page,
                chunk_index=chunk_index,
                start_char=start_char,
            )
            chunks.append(chunk)

        logger.info(f"Created {len(chunks)} chunks from document {document.title}")
        return chunks

    def _split_into_paragraphs(self, text: str) -> list[str]:
        """
        Split text into paragraphs based on newlines.

        Tries multiple strategies:
        1. Double newlines (standard paragraphs)
        2. Single newlines (if no double newlines found)
        3. Sentences (if very few paragraphs)
        """
        # Try double newlines first
        paragraphs = re.split(r"\n\n+", text)
        paragraphs = [p.strip() for p in paragraphs if p.strip()]

        # If we got very few paragraphs, try single newlines
        if len(paragraphs) < 3:
            paragraphs = re.split(r"\n+", text)
            paragraphs = [p.strip() for p in paragraphs if p.strip()]

        # If still very few, split by sentences
        if len(paragraphs) < 3:
            paragraphs = re.split(r"(?<=[.!?])\s+", text)
            paragraphs = [p.strip() for p in paragraphs if p.strip()]

        return paragraphs

    def _estimate_tokens(self, text: str) -> int:
        """
        Estimate token count (approximation: ~4 characters per token).

        More accurate would be to use a tokenizer, but this is a quick estimate.
        """
        return len(text) // 4

    def _get_overlap_text(self, text: str) -> str:
        """Get the last portion of text for overlap."""
        estimated_tokens = self._estimate_tokens(text)
        if estimated_tokens <= self.chunk_overlap:
            return text

        char_overlap = self.chunk_overlap * 4
        return text[-char_overlap:]

    def _create_chunk(
        self,
        document_id,
        text: str,
        page_number: int,
        chunk_index: int,
        start_char: int,
    ) -> Chunk:
        """Create a Chunk object."""
        end_char = start_char + len(text)
        token_count = self._estimate_tokens(text)

        return Chunk(
            document_id=document_id,
            text=text,
            page_number=page_number,
            chunk_index=chunk_index,
            start_char=start_char,
            end_char=end_char,
            token_count=token_count,
        )

    def process_document(self, doc_path: Path) -> tuple[Document, list[Chunk]]:
        """
        Complete document processing: extract and chunk.

        Args:
            doc_path: Path to document file (PDF, TXT, or DOCX)

        Returns:
            Tuple of (Document, list of Chunks)
        """
        # Use appropriate extraction method based on file type
        suffix = doc_path.suffix.lower()
        if suffix == ".pdf":
            document = self.extract_text_from_pdf(doc_path)
        elif suffix == ".txt":
            document = self.extract_text_from_txt(doc_path)
        elif suffix == ".docx":
            document = self.extract_text_from_docx(doc_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        chunks = self.chunk_document(document)
        document.processed = True

        logger.info(f"Document processing complete: {document.title} -> {len(chunks)} chunks")
        return document, chunks
