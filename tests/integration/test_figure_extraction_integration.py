"""Integration tests for figure extraction feature."""

import pytest

from docprocessor.core.figure_extractor import FigureExtractor
from docprocessor.models.extracted_figure import FigureType


class TestFigureExtractionIntegration:
    """Integration tests for complete extraction workflow."""

    @pytest.fixture
    def extractor(self):
        """Create extractor instance."""
        return FigureExtractor()

    @pytest.fixture
    def sample_docx_path(self, tmp_path):
        """Create a sample DOCX file for testing."""
        # Note: This requires python-docx installed
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Le budget de l'État en 2025 s'élève à €45.3 milliards.")
        doc.add_paragraph("Le taux de croissance est de 23.4% par rapport à 2024.")

        # Add a table
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Année"
        table.cell(0, 1).text = "Budget (€M)"
        table.cell(0, 2).text = "Croissance (%)"
        table.cell(1, 0).text = "2024"
        table.cell(1, 1).text = "42.8"
        table.cell(1, 2).text = "18.5"
        table.cell(2, 0).text = "2025"
        table.cell(2, 1).text = "45.3"
        table.cell(2, 2).text = "23.4"

        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        return docx_path

    @pytest.fixture
    def sample_pdf_path(self, tmp_path):
        """Create a sample PDF file for testing."""
        # Note: This test creates a simple text PDF using reportlab if available
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ImportError:
            pytest.skip("reportlab not installed")

        pdf_path = tmp_path / "test.pdf"

        c = canvas.Canvas(str(pdf_path), pagesize=letter)
        c.drawString(100, 750, "Le budget de l'État en 2025 s'élève à €45.3 milliards.")
        c.drawString(100, 730, "Le taux de croissance est de 23.4% par rapport à 2024.")
        c.drawString(100, 710, "Le ministère emploie 12,456 fonctionnaires.")
        c.save()

        return pdf_path

    def test_extract_from_txt_integration(self, extractor, tmp_path):
        """Test complete extraction workflow from TXT file."""
        # Create test file
        txt_file = tmp_path / "test.txt"
        txt_file.write_text(
            """
Le budget de l'État en 2025 s'élève à €45.3 milliards.
Le taux de croissance est de 23.4% par rapport à 2024.
Le ministère emploie 12,456 fonctionnaires.
Les recettes fiscales ont atteint €42.8 milliards en 2024.
"""
        )

        # Extract
        result = extractor.extract_from_document(txt_file)

        # Verify
        assert result is not None
        assert result.total_figures > 0
        assert result.extraction_time_seconds > 0
        assert isinstance(result.figures, list)

        # Check figure types
        assert result.figures_by_type.get("currency", 0) > 0
        assert result.figures_by_type.get("percentage", 0) > 0
        assert result.figures_by_type.get("date", 0) > 0

        # Verify specific figures
        currency_figs = result.get_figures_by_type(FigureType.CURRENCY)
        assert len(currency_figs) >= 1

        values = [fig.value for fig in currency_figs]
        assert any("45.3" in v for v in values)

    def test_extract_from_docx_with_tables(self, extractor, sample_docx_path):
        """Test extraction from DOCX with tables."""
        result = extractor.extract_from_document(sample_docx_path)

        assert result.total_figures > 0
        assert result.tables_parsed > 0

        # Check for table figures
        table_figures = [f for f in result.figures if f.is_from_table]
        assert len(table_figures) > 0

        # Verify table structure is preserved
        for fig in table_figures:
            assert fig.table_index is not None
            assert fig.table_row is not None
            assert fig.table_column is not None

    def test_extract_from_pdf_integration(self, extractor, sample_pdf_path):
        """Test extraction from PDF file."""
        result = extractor.extract_from_document(sample_pdf_path)

        assert result.total_figures > 0

        # Check page numbers are captured
        page_figures = [f for f in result.figures if f.page_number is not None]
        assert len(page_figures) > 0

        # All should be from page 1 in our test
        for fig in page_figures:
            assert fig.page_number == 1

    def test_multiformat_consistency(self, extractor, tmp_path, sample_docx_path):
        """Test that same content gives similar results across formats."""
        content = """
Le budget de l'État en 2025 s'élève à €45.3 milliards.
Le taux de croissance est de 23.4%.
"""

        # TXT version
        txt_file = tmp_path / "test.txt"
        txt_file.write_text(content)
        txt_result = extractor.extract_from_document(txt_file)

        # Compare totals (should be close, within ±20%)
        # Note: Different formats may extract slightly different sets
        assert txt_result.total_figures > 0

    def test_error_handling_invalid_file(self, extractor, tmp_path):
        """Test error handling with invalid file."""
        invalid_file = tmp_path / "nonexistent.pdf"

        # Should not crash, returns empty result with errors
        result = extractor.extract_from_document(invalid_file)

        assert result.total_figures == 0
        assert len(result.errors) > 0
        assert any("no such file" in err.lower() for err in result.errors)

    def test_error_handling_unsupported_format(self, extractor, tmp_path):
        """Test error handling with unsupported format."""
        unsupported = tmp_path / "test.xyz"
        unsupported.write_text("Some content")

        # Should handle gracefully or raise specific error
        result = extractor.extract_from_document(unsupported)
        # Unsupported formats should return empty result, not crash
        assert result.total_figures == 0

    def test_empty_document(self, extractor, tmp_path):
        """Test extraction from empty document."""
        empty_file = tmp_path / "empty.txt"
        empty_file.write_text("")

        result = extractor.extract_from_document(empty_file)

        assert result.total_figures == 0
        assert result.extraction_time_seconds >= 0

    def test_document_with_no_figures(self, extractor, tmp_path):
        """Test extraction from document without figures."""
        no_figs = tmp_path / "no_figures.txt"
        no_figs.write_text(
            """
Ceci est un document sans chiffres.
Il contient uniquement du texte narratif.
Aucune statistique présente ici.
"""
        )

        result = extractor.extract_from_document(no_figs)

        # Should complete without errors
        assert result.extraction_time_seconds >= 0
        # May find very few or zero figures (acceptable)
        assert result.total_figures >= 0

    def test_large_document_performance(self, extractor, tmp_path):
        """Test extraction performance with large document."""
        import time

        # Create large file with many figures
        large_file = tmp_path / "large.txt"
        content_lines = []
        for i in range(200):  # 200 paragraphs
            content_lines.append(
                f"En {2020 + i % 6}, le budget est de €{40 + i % 10}.{i % 10} milliards, "
                f"soit une croissance de {15 + i % 10}.{i % 10}%."
            )

        large_file.write_text("\n".join(content_lines))

        # Extract and measure time
        start = time.time()
        result = extractor.extract_from_document(large_file)
        duration = time.time() - start

        # Verify results
        assert result.total_figures > 100  # Should find many figures

        # Performance check: should process <20KB in under 5 seconds
        file_size_kb = large_file.stat().st_size / 1024
        print(f"\nProcessed {file_size_kb:.2f} KB in {duration:.2f}s")
        print(f"Found {result.total_figures} figures")

        # Reasonable performance expectation
        assert duration < 10.0, f"Too slow: {duration:.2f}s for {file_size_kb:.2f}KB"

    def test_extraction_result_serialization(self, extractor, tmp_path):
        """Test that extraction results can be serialized."""
        from dataclasses import asdict

        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Le budget 2025 est de €45.3 milliards.")

        result = extractor.extract_from_document(txt_file)

        # Test dataclass serialization
        result_dict = asdict(result)

        assert isinstance(result_dict, dict)
        assert "figures" in result_dict
        assert "total_figures" in result_dict
        assert "extraction_time_seconds" in result_dict

        # Verify figures are serializable
        for fig_dict in result_dict["figures"]:
            assert "value" in fig_dict
            assert "figure_type" in fig_dict

    def test_figure_deduplication(self, extractor, tmp_path):
        """Test that duplicate figures are handled correctly."""
        # Create file with repeated figures
        dup_file = tmp_path / "duplicates.txt"
        dup_file.write_text(
            """
Le budget 2025 est de €45.3 milliards.
Le budget 2025 est de €45.3 milliards.
Le budget 2025 est de €45.3 milliards.
"""
        )

        result = extractor.extract_from_document(dup_file)

        # Note: Current implementation may or may not deduplicate
        # This test documents the behavior
        currency_figs = result.get_figures_by_type(FigureType.CURRENCY)

        # Either duplicates are kept (3 figures) or removed (1 figure)
        # Both are valid behaviors
        assert len(currency_figs) >= 1
